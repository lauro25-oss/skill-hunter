import json
import re
from io import BytesIO

import pdfplumber
import docx
from groq import AsyncGroq
from app.config import settings

_groq: AsyncGroq | None = None


def _get_groq() -> AsyncGroq:
    global _groq
    if _groq is None:
        _groq = AsyncGroq(api_key=settings.groq_api_key)
    return _groq


# ── Extração de texto ────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text(x_tolerance=2, y_tolerance=2)
            if t:
                text_parts.append(t)
    return "\n\n".join(text_parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    doc = docx.Document(BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        text = _extract_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        text = _extract_docx(file_bytes)
    else:
        raise ValueError(f"Formato não suportado: '{ext}'. Use PDF ou DOCX.")

    if not text or len(text) < 50:
        raise ValueError(
            "Não foi possível extrair texto do arquivo. "
            "Verifique se o PDF não é uma imagem escaneada."
        )
    return text


# ── Prompt ───────────────────────────────────────────────────

SYSTEM_PROMPT = """Você é um extrator de dados de currículos especializado no mercado brasileiro.
Analise o texto e retorne SOMENTE um objeto JSON válido — sem markdown, sem explicações.

Regras:
- Use null para campos ausentes (nunca invente dados).
- "hard_skills" deve conter apenas habilidades técnicas concretas.
- "experiencias" deve listar CADA cargo com sua área e duração. Separe áreas distintas em itens diferentes.
  Exemplo: 3 anos em marketing e 2 anos como designer devem gerar dois itens separados.
- "anos_experiencia" deve ser a SOMA total de todos os anos listados em "experiencias".
- "resumo_ia" deve ter 3 a 5 linhas destacando pontos fortes do candidato.
- "localizacao" no formato "Cidade - UF" quando possível.

JSON de saída:
{
  "nome": "string",
  "email": "string | null",
  "telefone": "string | null",
  "localizacao": "string | null",
  "cargo_atual": "string | null",
  "anos_experiencia": "number | null",
  "experiencias": [
    {"cargo": "string", "area": "string", "anos": number}
  ],
  "hard_skills": ["string"],
  "resumo_ia": "string | null"
}"""


def _clean_json(raw: str) -> str:
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"```$", "", raw)
    return raw.strip()


# ── Pipeline principal ───────────────────────────────────────

async def parse_curriculum(file_bytes: bytes, filename: str) -> dict:
    """Extrai texto e usa Groq para estruturar os dados do currículo."""
    texto_bruto = extract_text(file_bytes, filename)

    parsed: dict = {
        "nome": filename.rsplit(".", 1)[0],
        "email": None,
        "telefone": None,
        "localizacao": None,
        "cargo_atual": None,
        "anos_experiencia": None,
        "experiencias": [],
        "hard_skills": [],
        "resumo_ia": None,
        "texto_bruto": texto_bruto,
    }

    if not settings.groq_api_key:
        return parsed

    client = _get_groq()

    response = await client.chat.completions.create(
        model=settings.groq_model,
        temperature=0,
        max_tokens=1024,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": texto_bruto[:8000]},
        ],
    )

    raw  = response.choices[0].message.content or "{}"
    raw  = _clean_json(raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        return parsed

    experiencias = _list_experiencias(data.get("experiencias"))
    anos_total   = _float(data.get("anos_experiencia"))
    # Se o Groq não somou, calcula a partir das experiências
    if anos_total is None and experiencias:
        anos_total = round(sum(e["anos"] for e in experiencias), 2)

    parsed.update({
        "nome":             _str(data.get("nome")) or parsed["nome"],
        "email":            _str(data.get("email")),
        "telefone":         _str(data.get("telefone")),
        "localizacao":      _str(data.get("localizacao")),
        "cargo_atual":      _str(data.get("cargo_atual")),
        "anos_experiencia": anos_total,
        "experiencias":     experiencias,
        "hard_skills":      _list_str(data.get("hard_skills")),
        "resumo_ia":        _str(data.get("resumo_ia")),
    })
    return parsed


def _str(v) -> str | None:
    if v is None or not str(v).strip():
        return None
    return str(v).strip()

def _float(v) -> float | None:
    try:
        return float(v)
    except (TypeError, ValueError):
        return None

def _list_str(v) -> list[str]:
    if not isinstance(v, list):
        return []
    return [str(i).strip() for i in v if str(i).strip()]


def _list_experiencias(v) -> list[dict]:
    if not isinstance(v, list):
        return []
    result = []
    for item in v:
        if not isinstance(item, dict):
            continue
        cargo = _str(item.get("cargo"))
        area  = _str(item.get("area"))
        anos  = _float(item.get("anos"))
        if cargo and anos is not None:
            result.append({"cargo": cargo, "area": area or cargo, "anos": anos})
    return result
